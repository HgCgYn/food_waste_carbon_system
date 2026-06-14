import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='標楷體', ascii_font='Times New Roman', size_pt=12, bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold

def set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0, first_line_indent=None):
    p.paragraph_format.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent

def main():
    doc = Document()
    
    # 設置紙張大小 (A4) 與邊界
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    # 讀取草稿
    lines = []
    for file in ['docs/report_draft.md', 'docs/report_draft_part2.md', 'docs/report_draft_part3.md']:
        with open(file, 'r', encoding='utf-8') as f:
            lines.extend(f.readlines())

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
            
        if line.startswith('# '):
            # Title
            p = doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=2.0)
            run = p.add_run(line.replace('# ', ''))
            set_font(run, size_pt=16, bold=True)
            
        elif i in [2, 3, 4] and '@' not in line and not line.startswith('【'):
            # Author names and affiliation (heuristic based on lines)
            p = doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
            run = p.add_run(line)
            set_font(run, size_pt=12, bold=False)
            
        elif '@' in line and not line.startswith('【'):
            # Email
            p = doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
            run = p.add_run(line)
            set_font(run, size_pt=12, bold=False)
            
        elif line.startswith('【摘要】'):
            p = doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
            run = p.add_run('【摘要】')
            set_font(run, size_pt=12, bold=True)
            
            p2 = doc.add_paragraph()
            set_paragraph_format(p2, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0)
            # The next line is the content of abstract
            content = line.replace('【摘要】', '').strip()
            if not content and i+1 < len(lines):
                i += 1
                content = lines[i].strip()
            run2 = p2.add_run(content)
            set_font(run2, size_pt=12, bold=False)
            
        elif line.startswith('【關鍵詞】'):
            p = doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
            run = p.add_run('【關鍵詞】')
            set_font(run, size_pt=12, bold=True)
            content = line.replace('【關鍵詞】', '').strip()
            run2 = p.add_run(content)
            set_font(run2, size_pt=12, bold=False)
            
        elif line[0].isdigit() and ' ' in line and line.split(' ')[0][-1] != '.':
            # Heading 1 or Heading 2 e.g., "1. 前言" or "1.1 研究背景"
            # In MD we wrote "1. 前言", "1.1 研究背景"
            p = doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=2.0)
            run = p.add_run(line)
            set_font(run, size_pt=14, bold=True)
            
        elif line.startswith('參考文獻'):
            p = doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=2.0)
            run = p.add_run(line)
            set_font(run, size_pt=14, bold=True)
            
        elif line.startswith('[IMAGE:'):
            img_path = line.replace('[IMAGE:', '').replace(']', '').strip()
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run()
                run.add_picture(img_path, width=Cm(15))
            else:
                print(f"Warning: Image {img_path} not found.")
                
        elif line.startswith('[TABLE_ACCURACY]'):
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '模型配置'
            hdr_cells[1].text = '總體準確率 (102個物件)'
            table.rows[1].cells[0].text = '純 YOLO (Baseline)'
            table.rows[1].cells[1].text = '26.5%'
            table.rows[2].cells[0].text = 'YOLO + Gemini 2.5 Flash'
            table.rows[2].cells[1].text = '42.2%'
            table.rows[3].cells[0].text = 'YOLO + GPT-4o'
            table.rows[3].cells[1].text = '44.1%'
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        set_font(paragraph.runs[0], size_pt=12)
            doc.add_paragraph() # Add some spacing
            
        elif line.startswith('[TABLE_CULTURE]'):
            table = doc.add_table(rows=3, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '飲食分類'
            hdr_cells[1].text = 'YOLO 原始'
            hdr_cells[2].text = 'YOLO + Gemini 2.5'
            hdr_cells[3].text = 'YOLO + GPT-4o'
            
            table.rows[1].cells[0].text = 'Eastern (中式/亞洲)'
            table.rows[1].cells[1].text = '34.0%'
            table.rows[1].cells[2].text = '44.0%'
            table.rows[1].cells[3].text = '42.0%'
            
            table.rows[2].cells[0].text = 'Western (西式)'
            table.rows[2].cells[1].text = '19.2%'
            table.rows[2].cells[2].text = '40.4%'
            table.rows[2].cells[3].text = '34.6%'
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.runs:
                            set_font(paragraph.runs[0], size_pt=12)
            doc.add_paragraph()

        elif line.startswith('[TABLE_MCNEMAR]'):
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '檢定組合'
            hdr_cells[1].text = 'p-value'
            hdr_cells[2].text = '顯著性 (α=0.05)'
            
            table.rows[1].cells[0].text = 'YOLO vs GPT-4o'
            table.rows[1].cells[1].text = '0.0247'
            table.rows[1].cells[2].text = '顯著 (*)'
            
            table.rows[2].cells[0].text = 'YOLO vs Gemini 2.5 Flash'
            table.rows[2].cells[1].text = '0.0033'
            table.rows[2].cells[2].text = '極顯著 (**)'
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.runs:
                            set_font(paragraph.runs[0], size_pt=12)
            doc.add_paragraph()
            
        else:
            # Normal text
            p = doc.add_paragraph()
            # 內縮 0.63 公分 (0.25 inch)
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0, first_line_indent=Cm(0.63))
            run = p.add_run(line)
            set_font(run, size_pt=12, bold=False)
            
        i += 1

    doc.save('專題報告_自動排版_v11.docx')
    print("Report generated successfully: 專題報告_自動排版_v11.docx")

if __name__ == '__main__':
    main()
